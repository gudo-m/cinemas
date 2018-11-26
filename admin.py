from docx import Document
from docx.shared import Inches
from pptx import Presentation
import sql_worker
import colorama
import os
import datetime
import time
import xlsxwriter

if 'PYCHARM_HOSTED' in os.environ:
    convert = False  # in PyCharm, we should disable convert
    strip = False
    print("Hi! You are using PyCharm")
else:
    convert = None
    strip = None

colorama.init(convert=convert, strip=strip)


def loop(request):
    def handle_this(hundeling_dict):
        commas = ("'", '"')
        for handle_this_id in hundeling_dict:
            if answer[handle_this_id][-1] not in commas:
                print(uncorrect_answer)
                continue
        data = {}
        for hndl_id in hundeling_dict:
            try:
                index_of_start = answer[hndl_id].index(hundeling_dict[hndl_id][0]) \
                                 + len(hundeling_dict[hndl_id][0])
            except Exception:
                try:
                    index_of_start = answer[hndl_id].index(hundeling_dict[hndl_id][1]) \
                                     + len(hundeling_dict[hndl_id][1])
                except Exception as e:
                    print(e)
                continue
            data[hundeling_dict[hndl_id][2]] = answer[hndl_id][index_of_start:-1]
        return data
    answer = ''
    print('Здравствуйте, {}!'.format(request.user['login']))
    uncorrect_answer = colorama.Fore.RED + '\n\n\n\nВведите корректную команду!' + colorama.Fore.RESET
    while " ".join(answer).lower() != 'exit':
        print(
            '\n\nСписок команд:\n\n'
            '"CREATE CINEMA name=\'<name>\' address=\'<address>\'" - Создать кинотеатр с названием и адресом\n'
            '"CREATE CINEMA_HALL cinema=\'<cinema_name>\' name=\'<name>\' num_places=\'<num_places>\' num_cols=\'<num_cols>\'"'
            '- Создать кинозал с названием и <num_places> количеством мест в кинотеатре с названием <cinema_name>\n'
            '"ADD_FILM halls=\'<cinema_hall_name1>-<cinema_hall_name2>\' title=\'<title>\' actual=\'<num_days>\'" - Добавить фильм с названием'
            '<title> во все кинозалы, перечисленные через тире в "halls", фильм будет доступен <num_days> дней\n'
            '"SHOW CINEMAS" - Показать все кинотеатры\n'
            '"SHOW CINEMAS_HALLS" - Показать все кинозалы\n'
            '"SHOW FILMS" - Показать текущие фильмы\n'
            '"COLLECT_STATISTIC" - Экспортировать статистику\n'
            '"DELETE CINEMA name=\'<name>\'" - удаляет кинотеатр\n'
            '"DELETE CINEMA_HALL name=\'<name>\'" - удаляет кинозал\n'
            '"CHANGE_PASSWORD password=\'<password>\'" - изменяет пароль\n'
            '"CREATE_SUPERUSER login=\'<login>\' password=\'<password>\'" - создать аккаунт администратора\n'
            '"EXIT" - выйти из программы\nВнимание! Все пробелы внутри кавычек заменяйте на "_", система не сможет распознать пробел\n\n{}@localmachine:~$ '.format(request.user['login']),
            end=''
        )
        answer = input().split()
        print()
        print()
        print()

        if len(answer) < 1:
            print(uncorrect_answer)
            continue
        elif answer[0] == 'CREATE':
            if len(answer) not in (4, 6):
                print(uncorrect_answer)
                continue
            elif answer[1] == 'CINEMA':
                if len(answer) != 4:
                    print(uncorrect_answer)
                    continue
                else:
                    data = handle_this({2: ('name=\'', 'name="', 'name'), 3: ('address=\'', 'address="', 'addr')})
                    cinema_with_this_name = request.cursor.execute('SELECT * FROM cinemas WHERE name="{}"'.format(data['name']))
                    if len(cinema_with_this_name) > 0:
                        print(colorama.Fore.RED + 'Кинотеатр с таким названием уже существует!' + colorama.Fore.RESET)
                        continue
                    request.cursor.execute('INSERT INTO cinemas (name, address) VALUES ("{}", "{}")'.format(
                        data['name'],
                        data['addr'])
                    )
                print('Успешно!')
            elif answer[1] == 'CINEMA_HALL':
                if len(answer) != 6:
                    print(uncorrect_answer)
                    continue
                else:
                    data = handle_this({
                        2: ('cinema=\'', 'cinema="', 'cinema'),
                        3: ('name=\'', 'name="', 'name'),
                        4: ('num_places=\'', 'num_places="', 'num_places'),
                        5: ('num_cols=\'', 'num_cols="', 'num_cols'),
                    })
                    cinemas = request.cursor.execute('SELECT id FROM cinemas WHERE name="{}"'.format(data['cinema']))
                    cinema_hall_with_this_name = request.cursor.execute('SELECT * FROM cinemas_halls WHERE name="{}"'.format(data['name']))
                    if len(cinema_hall_with_this_name) > 0:
                        print(colorama.Fore.RED + 'Кинозал с таким названием уже существует!' + colorama.Fore.RESET)
                        continue
                    if len(cinemas) < 1:
                        print(colorama.Fore.RED + 'Кинотеатр с таким названием не найден!' + colorama.Fore.RESET)
                        continue
                    data['cinema_id'] = cinemas[0][0]
                    if not data['num_places'].isdigit():
                        print(uncorrect_answer)
                        continue
                    data['places'] = "".join(['|{}-0'.format(i) for i in range(int(data['num_places']))])
                    request.cursor.execute(
                        'INSERT INTO cinemas_halls (name, places, cols, cinema_id) VALUES ("{}", "{}", "{}", "{}")'.format(
                            data['name'],
                            data['places'],
                            data['num_cols'],
                            data['cinema_id']
                        )
                    )
                print('Успешно!')
            else:
                print(uncorrect_answer)
                continue
        elif answer[0] == 'DELETE':
                if len(answer) != 3:
                    print(uncorrect_answer)
                    continue
                elif answer[1] == 'CINEMA':
                    data = handle_this({
                        2: ('name=\'', 'name="', 'name'),
                    })
                    print('Найдено совпадений: {}'.format(len(num_found)))
                    if len(num_found) > 0:
                        request.cursor.execute('DELETE FROM cinemas WHERE name="{}"'.format(data['name']))
                        print('Успешно удалено!')
                elif answer[1] == 'CINEMA_HALL':
                    data = handle_this({
                        2: ('name=\'', 'name="', 'name'),
                    })
                    num_found = request.cursor.execute('SELECT * FROM cinemas_halls WHERE name="{}"'.format(data['name']))
                    print('Найдено совпадений: {}'.format(len(num_found)))
                    if len(num_found) > 0:
                        request.cursor.execute('DELETE FROM cinemas_halls WHERE name="{}"'.format(data['name']))
                        print('Успешно удалено!')
                else:
                    print(uncorrect_answer)
                    continue
        elif answer[0] == 'SHOW':
            if len(answer) != 2:
                print(uncorrect_answer)
                continue
            elif answer[1] == 'CINEMAS':
                cinemas = request.cursor.execute('SELECT * FROM cinemas')
                if len(cinemas) > 0:
                    print('№) <name> - <address>')
                else:
                    print(colorama.Fore.RED + 'Не найдено ни одного кинотеатра' + colorama.Fore.RESET)
                for cinema in cinemas:
                    print('{}) {} - {}'.format(cinema[0], cinema[1], cinema[2]))
            elif answer[1] == 'CINEMAS_HALLS':
                cinemas_halls = request.cursor.execute('SELECT * FROM cinemas_halls')
                if len(cinemas_halls) > 0:
                    print('№) <name> - <cinema_name>')
                else:
                    print(colorama.Fore.RED + 'Не найдено ни одного кинозала' + colorama.Fore.RESET)
                for hall in cinemas_halls:
                    print('{}) {} - {}'.format(
                        hall[0],
                        hall[1],
                        request.cursor.execute('SELECT name FROM cinemas WHERE id="{}"'.format(hall[4]))[0][0]
                    ))
                    places = hall[2].split('|')[1:]

                    len_of_place = len(str(len(places)-1))
                    num_cols = int(hall[3])
                    num_spaces_in_row = num_cols - 1
                    len_of_row = len_of_place * num_cols + num_spaces_in_row

                    print()
                    print(colorama.Fore.YELLOW + 'Экран'.rjust(len_of_row//2+5-3, '-').ljust(len_of_row, '-') + colorama.Fore.RESET)
                    for place_id in range(len(places)):
                        place = places[place_id].split('-')
                        if place[1] == '1':
                            print(colorama.Fore.RED + place[0].rjust(len(str(len(places)-1)), '0') + colorama.Fore.RESET, end=' ')
                        if place[1] == '0':
                            print(colorama.Fore.GREEN + place[0].rjust(len(str(len(places)-1)), '0') + colorama.Fore.RESET, end=' ')
                        if (place_id+1) % int(hall[3]) == 0:
                            print()
                    print()
            elif answer[1] == 'FILMS':
                print('№) <title> - <halls>')
                now = datetime.datetime.now().timestamp()
                films = request.cursor.execute('SELECT * FROM films WHERE actual_to > {}'.format(now))
                for film in films:
                    print('{}) {} - {}'.format(film[0], film[1], ", ".join([hall.split('[')[0] for hall in film[2].split()])))
            else:
                print(uncorrect_answer)
                continue
        elif answer[0] == 'EXIT' and len(answer) == 1:
            return None
        elif answer[0] == 'COLLECT_STATISTIC' and len(answer) == 1:
            document = Document()

            document.add_heading('Статистика CinemasControlSystem', 0)

            document.add_paragraph('Фильмы и занятость залов')
            document.add_paragraph('("Б" - означает забронированное место)')
            films = request.cursor.execute('SELECT * FROM films')
            for film in films:
                document.add_paragraph('\n')
                document.add_paragraph('Название: {}'.format(film[1])).bold = True
                document.add_paragraph('\n')
                for hall in film[2].split():
                    hall_name = hall.split('[')[0]
                    hall_from_bd = request.cursor.execute('SELECT * FROM cinemas_halls WHERE name="{}"'.format(hall_name))[0]
                    places = hall.split('[')[1][:-1].split('|')[1:]

                    len_of_place = len(str(len(places) - 1))
                    num_cols = int(hall_from_bd[3])
                    num_spaces_in_row = num_cols - 1
                    len_of_row = len_of_place * num_cols + num_spaces_in_row

                    document.add_paragraph('\n')
                    p = document.add_paragraph('Экран'.rjust(len_of_row // 2 + 5 - 3, '-').ljust(len_of_row, '-'))
                    p.add_run('\n')
                    for place_id in range(len(places)):
                        place = places[place_id].split('-')
                        if place[1] == '1':
                            p.add_run('Б ')
                        if place[1] == '0':
                            p.add_run(place[0].rjust(len(str(len(places) - 1)), '0') + ' ')
                        if (place_id + 1) % int(hall_from_bd[3]) == 0:
                            p.add_run('\n')
                    p.add_run('\n')
                document.add_paragraph('\n')
                document.add_paragraph('\n')
                document.add_paragraph('\n')

            document.save('statistic.docx')
            print('Успешно сохранено в файл с названием "statistic.docx"!')
        elif answer[0] == 'ADD_FILM' and len(answer) == 4:
            data = handle_this({1: ('halls=\'', 'halls="', 'halls'), 2: ('title=\'', 'title="', 'title'), 3: ('actual=\'', 'actual="', 'actual')})
            films_with_this_title = request.cursor.execute('SELECT id FROM films WHERE title="{}"'.format(data['title']))
            if len(films_with_this_title) > 0:
                print(colorama.Fore.RED + 'Фильм с таким названием существует!' + colorama.Fore.RESET)
                continue
            actual_to = datetime.datetime.now() + datetime.timedelta(days=int(data['actual']))
            flag = True
            for hall_name in data['halls'].split('-'):
                hall_with_this_name = request.cursor.execute('SELECT id FROM cinemas_halls WHERE name="{}"'.format(hall_name))
                if len(hall_with_this_name) == 0:
                    print(colorama.Fore.RED + 'Зала с таким названием не существует!' + colorama.Fore.RESET)
                    flag = False
                    break
            if not flag:
                continue
            else:
                actual_to = actual_to.timestamp()
                halls = ''
                for hall in data['halls'].split('-'):
                    halls += '{}[{}] '.format(hall, request.cursor.execute('SELECT places FROM cinemas_halls WHERE name="{}"'.format(hall))[0][0])
                request.cursor.execute('INSERT INTO films (title, halls, actual_to) VALUES ("{}", "{}", "{}")'.format(data['title'], halls, actual_to))
                print('Успешно!')
        elif answer[0] == 'CHANGE_PASSWORD' and len(answer) == 2:
            data = handle_this({1: ('password=\'', 'password="', 'password')})
            request.cursor.execute('UPDATE auth SET password="{}" WHERE login="{}"'.format(data['password'], request.user['login']))
            print('Успешно!')
        elif answer[0] == 'CREATE_SUPERUSER' and len(answer) == 3:
            data = handle_this({1: ('login=\'', 'login="', 'login'), 2: ('password=\'', 'password="', 'password')})
            users_with_this_login = request.cursor.execute('SELECT * FROM auth WHERE login="{}"'.format(data['login']))
            if len(users_with_this_login) > 0:
                print(colorama.Fore.RED + 'Аккаунта с таким логином существует!' + colorama.Fore.RESET)
                continue
            request.cursor.execute('INSERT INTO auth (login, password, is_admin) VALUES ("{}", "{}", "1")'.format(data['login'], data['password']))
            print('Успешно!')
        else:
            print(uncorrect_answer)
            continue


if __name__ == '__main__':
    login = input('Введите логин: ')
    password = input('Введите пароль: ')
    request = sql_worker.Request()
    request_answer = request.auth_admin(login, password)
    request_answer = request_answer.get('success', request_answer.get('error', None))
    if request_answer is not None:
        print(request_answer)
    if request.user['login'] != 'Anonym':
        loop(request)
    input('Нажмите ENTER, чтобы завершить программу')
    request.kill()
